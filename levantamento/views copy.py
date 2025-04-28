from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Projeto, Beneficiario, Confrontante, Vertice
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from django.http import HttpResponse
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required


# View de login
def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('index')  # Redireciona para a view de geração de PDF
        else:
            messages.error(request, 'Usuário ou senha inválidos.')
            return render(request, 'levantamento/login.html')
    return render(request, 'levantamento/login.html')

# View de logout
def logout_view(request):
    logout(request)
    return redirect('login')

# Exemplo: Proteger a view de geração de PDF
@login_required
def gerar_pdf(request, projeto_id):
    # Seu código existente para geração de PDF
    # Exemplo: apenas usuários autenticados podem acessar essa view
    projeto = Projeto.objects.get(id=projeto_id)
    beneficiarios = projeto.beneficiarios.all()
    confrontantes = projeto.confrontantes.all()
    
    # Seu código de geração de PDF aqui...
    elements = []
    # (Código anterior para gerar o PDF, como a tabela de assinaturas, etc.)
    
    return render(request, 'levantamento/index.html', {'elements': elements})

def index(request):
    projetos = Projeto.objects.all()
    beneficiarios = []
    confrontantes = []
    vertices = []
    projeto_selecionado = None

    if 'projeto_selecionado_id' in request.session:
        try:
            projeto_selecionado = Projeto.objects.get(id=request.session['projeto_selecionado_id'])
        except Projeto.DoesNotExist:
            projeto_selecionado = projetos.order_by('-id').first()
    else:
        projeto_selecionado = projetos.order_by('-id').first()

    if projeto_selecionado:
        beneficiarios = Beneficiario.objects.filter(projeto=projeto_selecionado)
        confrontantes = Confrontante.objects.filter(projeto=projeto_selecionado)
        vertices = Vertice.objects.filter(projeto=projeto_selecionado)

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'selecionar_projeto':
            projeto_id = request.POST.get('projeto_filtro')
            try:
                projeto_selecionado = Projeto.objects.get(id=projeto_id)
                request.session['projeto_selecionado_id'] = projeto_id
                beneficiarios = Beneficiario.objects.filter(projeto=projeto_selecionado)
                confrontantes = Confrontante.objects.filter(projeto=projeto_selecionado)
                vertices = Vertice.objects.filter(projeto=projeto_selecionado)
            except Projeto.DoesNotExist:
                messages.error(request, 'Projeto selecionado não existe.')

        elif action == 'add_projeto':
            nome = request.POST.get('nome_projeto')
            endereco = request.POST.get('endereco_projeto')
            area = request.POST.get('area_projeto')
            perimetro = request.POST.get('perimetro_projeto')
            epoca_medicao = request.POST.get('epoca_medicao')
            instrumento = request.POST.get('instrumento')
            try:
                projeto = Projeto.objects.create(
                    nome=nome,
                    endereco=endereco,
                    area=float(area),
                    perimetro=float(perimetro),
                    epoca_medicao=epoca_medicao,
                    instrumento=instrumento
                )
                request.session['projeto_selecionado_id'] = projeto.id
                messages.success(request, 'Projeto adicionado com sucesso!')
            except ValueError as e:
                messages.error(request, f'Erro ao adicionar projeto: {str(e)}')

        elif action == 'add_beneficiario':
            projeto_id = request.POST.get('projeto_ben')
            nome = request.POST.get('nome_ben')
            cpf_cnpj = request.POST.get('cpf_cnpj_ben')
            rua = request.POST.get('rua_ben')
            numero = request.POST.get('numero_ben')
            bairro = request.POST.get('bairro_ben')
            cidade = request.POST.get('cidade_ben')
            try:
                projeto = Projeto.objects.get(id=projeto_id)
                Beneficiario.objects.create(
                    projeto=projeto,
                    nome=nome,
                    cpf_cnpj=cpf_cnpj,
                    rua=rua,
                    numero=numero,
                    bairro=bairro,
                    cidade=cidade
                )
                messages.success(request, 'Beneficiário adicionado com sucesso!')
            except Projeto.DoesNotExist:
                messages.error(request, 'Projeto selecionado não existe.')
            except Exception as e:
                messages.error(request, f'Erro ao adicionar beneficiário: {str(e)}')

        elif action == 'importar_beneficiarios':
            projeto_id = request.POST.get('projeto_ben')
            arquivo = request.FILES.get('arquivo_beneficiarios')
            
            if projeto_id and arquivo:
                try:
                    projeto = Projeto.objects.get(id=projeto_id)
                    try:
                        conteudo = arquivo.read().decode('utf-8').splitlines()
                    except UnicodeDecodeError:
                        arquivo.seek(0)
                        try:
                            conteudo = arquivo.read().decode('latin-1').splitlines()
                        except UnicodeDecodeError:
                            arquivo.seek(0)
                            conteudo = arquivo.read().decode('windows-1252').splitlines()
                    
                    for linha in conteudo:
                        campos = linha.strip().split('\t')
                        if len(campos) != 6:
                            messages.error(request, f'Formato inválido na linha: {linha}')
                            continue
                        nome, cpf_cnpj, rua, numero, bairro, cidade = campos
                        
                        Beneficiario.objects.create(
                            projeto=projeto,
                            nome=nome,
                            cpf_cnpj=cpf_cnpj,
                            rua=rua,
                            numero=numero,
                            bairro=bairro,
                            cidade=cidade
                        )
                    messages.success(request, 'Beneficiários importados com sucesso!')
                except Projeto.DoesNotExist:
                    messages.error(request, 'Projeto selecionado não existe.')
                except Exception as e:
                    messages.error(request, f'Erro ao importar beneficiários: {str(e)}')
            else:
                messages.error(request, 'Selecione um projeto e um arquivo TXT.')

        elif action == 'edit_beneficiario':
            beneficiario_id = request.POST.get('beneficiario_id')
            nome = request.POST.get('nome_ben')
            cpf_cnpj = request.POST.get('cpf_cnpj_ben')
            rua = request.POST.get('rua_ben')
            numero = request.POST.get('numero_ben')
            bairro = request.POST.get('bairro_ben')
            cidade = request.POST.get('cidade_ben')
            try:
                beneficiario = Beneficiario.objects.get(id=beneficiario_id)
                beneficiario.nome = nome
                beneficiario.cpf_cnpj = cpf_cnpj
                beneficiario.rua = rua
                beneficiario.numero = numero
                beneficiario.bairro = bairro
                beneficiario.cidade = cidade
                beneficiario.save()
                messages.success(request, 'Beneficiário atualizado com sucesso!')
            except Beneficiario.DoesNotExist:
                messages.error(request, 'Beneficiário não encontrado.')
            except Exception as e:
                messages.error(request, f'Erro ao atualizar beneficiário: {str(e)}')

        elif action == 'delete_beneficiario':
            beneficiario_id = request.POST.get('beneficiario_id')
            try:
                beneficiario = Beneficiario.objects.get(id=beneficiario_id)
                beneficiario.delete()
                messages.success(request, 'Beneficiário excluído com sucesso!')
            except Beneficiario.DoesNotExist:
                messages.error(request, 'Beneficiário não encontrado.')
            except Exception as e:
                messages.error(request, f'Erro ao excluir beneficiário: {str(e)}')

        elif action == 'add_confrontante':
            projeto_id = request.POST.get('projeto_con')
            nome = request.POST.get('nome_con')
            cpf_cnpj = request.POST.get('cpf_cnpj_con')
            direcao = request.POST.get('direcao_con')
            rua = request.POST.get('rua_con')
            numero = request.POST.get('numero_con')
            bairro = request.POST.get('bairro_con')
            cidade = request.POST.get('cidade_con')
            try:
                projeto = Projeto.objects.get(id=projeto_id)
                Confrontante.objects.create(
                    projeto=projeto,
                    nome=nome,
                    cpf_cnpj=cpf_cnpj,
                    direcao=direcao,
                    rua=rua,
                    numero=numero,
                    bairro=bairro,
                    cidade=cidade
                )
                messages.success(request, 'Confrontante adicionado com sucesso!')
            except Projeto.DoesNotExist:
                messages.error(request, 'Projeto selecionado não existe.')
            except Exception as e:
                messages.error(request, f'Erro ao adicionar confrontante: {str(e)}')

        elif action == 'edit_confrontante':
            confrontante_id = request.POST.get('confrontante_id')
            nome = request.POST.get('nome_con')
            cpf_cnpj = request.POST.get('cpf_cnpj_con')
            direcao = request.POST.get('direcao_con')
            rua = request.POST.get('rua_con')
            numero = request.POST.get('numero_con')
            bairro = request.POST.get('bairro_con')
            cidade = request.POST.get('cidade_con')
            try:
                confrontante = Confrontante.objects.get(id=confrontante_id)
                confrontante.nome = nome
                confrontante.cpf_cnpj = cpf_cnpj
                confrontante.direcao = direcao
                confrontante.rua = rua
                confrontante.numero = numero
                confrontante.bairro = bairro
                confrontante.cidade = cidade
                confrontante.save()
                messages.success(request, 'Confrontante atualizado com sucesso!')
            except Confrontante.DoesNotExist:
                messages.error(request, 'Confrontante não encontrado.')
            except Exception as e:
                messages.error(request, f'Erro ao atualizar confrontante: {str(e)}')

        elif action == 'delete_confrontante':
            confrontante_id = request.POST.get('confrontante_id')
            try:
                confrontante = Confrontante.objects.get(id=confrontante_id)
                confrontante.delete()
                messages.success(request, 'Confrontante excluído com sucesso!')
            except Confrontante.DoesNotExist:
                messages.error(request, 'Confrontante não encontrado.')
            except Exception as e:
                messages.error(request, f'Erro ao excluir confrontante: {str(e)}')

        elif action == 'toggle_confrontante_pdf':
            excluir_ids = request.POST.getlist('excluir_confrontantes')  # Lista de IDs dos confrontantes a excluir
            confrontantes = Confrontante.objects.filter(projeto__id=request.POST.get('projeto_filtro'))
            for confrontante in confrontantes:
                # Se o ID do confrontante está na lista de exclusão, marca como True, senão False
                confrontante.excluir_do_pdf = str(confrontante.id) in excluir_ids
                confrontante.save()
            messages.success(request, 'Seleção de confrontantes atualizada!')
            return redirect('index')

        elif action == 'importar_confrontantes':
            projeto_id = request.POST.get('projeto_con')
            arquivo = request.FILES.get('arquivo_confrontantes')
            
            if projeto_id and arquivo:
                try:
                    projeto = Projeto.objects.get(id=projeto_id)
                    try:
                        conteudo = arquivo.read().decode('utf-8').splitlines()
                    except UnicodeDecodeError:
                        arquivo.seek(0)
                        try:
                            conteudo = arquivo.read().decode('latin-1').splitlines()
                        except UnicodeDecodeError:
                            arquivo.seek(0)
                            conteudo = arquivo.read().decode('windows-1252').splitlines()
                    
                    for linha in conteudo:
                        campos = linha.strip().split('\t')
                        if len(campos) != 7:
                            messages.error(request, f'Formato inválido na linha: {linha}')
                            continue
                        nome, cpf_cnpj, direcao, rua, numero, bairro, cidade = campos
                        if direcao not in ['Frente', 'Fundos', 'Direito', 'Esquerdo']:
                            messages.error(request, f'Direção inválida na linha: {linha}. Use Frente, Fundos, Direito ou Esquerdo.')
                            continue
                        
                        Confrontante.objects.create(
                            projeto=projeto,
                            nome=nome,
                            cpf_cnpj=cpf_cnpj,
                            direcao=direcao,
                            rua=rua,
                            numero=numero,
                            bairro=bairro,
                            cidade=cidade
                        )
                    messages.success(request, 'Confrontantes importados com sucesso!')
                except Projeto.DoesNotExist:
                    messages.error(request, 'Projeto selecionado não existe.')
                except Exception as e:
                    messages.error(request, f'Erro ao importar confrontantes: {str(e)}')
            else:
                messages.error(request, 'Selecione um projeto e um arquivo TXT.')

        elif action == 'importar_vertices':
            projeto_id = request.POST.get('projeto_ver')
            arquivo = request.FILES.get('arquivo_vertices')
            
            if projeto_id and arquivo:
                try:
                    projeto = Projeto.objects.get(id=projeto_id)
                    try:
                        conteudo = arquivo.read().decode('utf-8').splitlines()
                    except UnicodeDecodeError:
                        arquivo.seek(0)
                        try:
                            conteudo = arquivo.read().decode('latin-1').splitlines()
                        except UnicodeDecodeError:
                            arquivo.seek(0)
                            conteudo = arquivo.read().decode('windows-1252').splitlines()
                    
                    for linha in conteudo:
                        campos = linha.strip().split('\t')
                        if len(campos) < 6:
                            messages.error(request, f'Formato inválido na linha: {linha}')
                            continue
                        de_vertice, para_vertice, longitude, latitude, distancia, confrontante_nome = campos[:6]
                        confrontante_cpf_cnpj = campos[6] if len(campos) > 6 else ''
                        
                        vertice_data = {
                            'projeto': projeto,
                            'de_vertice': de_vertice,
                            'para_vertice': para_vertice,
                            'longitude': longitude,
                            'latitude': latitude,
                            'distancia': float(distancia),
                            'confrontante_texto': confrontante_nome
                        }
                        if confrontante_cpf_cnpj:
                            try:
                                confrontante = Confrontante.objects.get(cpf_cnpj=confrontante_cpf_cnpj, projeto=projeto)
                                vertice_data['confrontante'] = confrontante
                                vertice_data['confrontante_texto'] = ''
                            except Confrontante.DoesNotExist:
                                pass
                        Vertice.objects.create(**vertice_data)
                    messages.success(request, 'Vértices importados com sucesso!')
                except Projeto.DoesNotExist:
                    messages.error(request, 'Projeto selecionado não existe.')
                except ValueError as e:
                    messages.error(request, f'Erro ao importar vértices: {str(e)}')
                except Exception as e:
                    messages.error(request, f'Erro inesperado: {str(e)}')
            else:
                messages.error(request, 'Selecione um projeto e um arquivo TXT.')

        elif action == 'add_vertice':
            projeto_id = request.POST.get('projeto_ver')
            de_vertice = request.POST.get('de_vertice')
            para_vertice = request.POST.get('para_vertice')
            longitude = request.POST.get('longitude_ver')
            latitude = request.POST.get('latitude_ver')
            distancia = request.POST.get('distancia_ver')
            confrontante_id = request.POST.get('confrontante_ver')
            confrontante_texto = request.POST.get('confrontante_texto')
            try:
                projeto = Projeto.objects.get(id=projeto_id)
                distancia = float(distancia) if distancia else 0.0
                vertice_data = {
                    'projeto': projeto,
                    'de_vertice': de_vertice,
                    'para_vertice': para_vertice,
                    'longitude': longitude,
                    'latitude': latitude,
                    'distancia': distancia,
                    'confrontante_texto': confrontante_texto
                }
                if confrontante_id:
                    confrontante = Confrontante.objects.get(id=confrontante_id, projeto=projeto)
                    vertice_data['confrontante'] = confrontante
                    vertice_data['confrontante_texto'] = ''
                Vertice.objects.create(**vertice_data)
                messages.success(request, 'Vértice adicionado com sucesso!')
            except Projeto.DoesNotExist:
                messages.error(request, 'Projeto selecionado não existe.')
            except Confrontante.DoesNotExist:
                messages.error(request, 'Confrontante selecionado não existe.')
            except ValueError as e:
                messages.error(request, f'Erro ao adicionar vértice: {str(e)}')
            except Exception as e:
                messages.error(request, f'Erro inesperado: {str(e)}')

        elif action == 'edit_vertice':
            vertice_id = request.POST.get('vertice_id')
            de_vertice = request.POST.get('de_vertice')
            para_vertice = request.POST.get('para_vertice')
            longitude = request.POST.get('longitude_ver')
            latitude = request.POST.get('latitude_ver')
            distancia = request.POST.get('distancia_ver')
            confrontante_id = request.POST.get('confrontante_ver')
            confrontante_texto = request.POST.get('confrontante_texto')
            try:
                vertice = Vertice.objects.get(id=vertice_id)
                distancia = float(distancia) if distancia else 0.0
                vertice.de_vertice = de_vertice
                vertice.para_vertice = para_vertice
                vertice.longitude = longitude
                vertice.latitude = latitude
                vertice.distancia = distancia
                vertice.confrontante_texto = confrontante_texto
                if confrontante_id:
                    confrontante = Confrontante.objects.get(id=confrontante_id)
                    vertice.confrontante = confrontante
                    vertice.confrontante_texto = ''
                else:
                    vertice.confrontante = None
                vertice.save()
                messages.success(request, 'Vértice atualizado com sucesso!')
            except Vertice.DoesNotExist:
                messages.error(request, 'Vértice não encontrado.')
            except Confrontante.DoesNotExist:
                messages.error(request, 'Confrontante selecionado não existe.')
            except ValueError as e:
                messages.error(request, f'Erro ao atualizar vértice: {str(e)}')
            except Exception as e:
                messages.error(request, f'Erro inesperado: {str(e)}')

        elif action == 'delete_vertice':
            vertice_id = request.POST.get('vertice_id')
            try:
                vertice = Vertice.objects.get(id=vertice_id)
                vertice.delete()
                messages.success(request, 'Vértice excluído com sucesso!')
            except Vertice.DoesNotExist:
                messages.error(request, 'Vértice não encontrado.')
            except Exception as e:
                messages.error(request, f'Erro ao excluir vértice: {str(e)}')

        elif action == 'gerar_memorial':
            projeto_id = request.POST.get('projeto_memorial')
            try:
                projeto = Projeto.objects.get(id=projeto_id)
                vertices = Vertice.objects.filter(projeto=projeto)
                beneficiarios = Beneficiario.objects.filter(projeto=projeto)
                confrontantes = Confrontante.objects.filter(projeto=projeto)

                # Log para depuração
                print(f"Projeto ID: {projeto_id}")
                print(f"Beneficiários encontrados: {len(beneficiarios)}")
                print(f"Confrontantes encontrados: {len(confrontantes)}")
                print(f"Vértices encontrados: {len(vertices)}")

                doc = Document()

                # Definir margens (2,5 cm em todos os lados)
                for section in doc.sections:
                    section.left_margin = Cm(2.5)
                    section.right_margin = Cm(2.5)
                    section.top_margin = Cm(2.5)
                    section.bottom_margin = Cm(2.5)

                # Definir fonte padrão para o documento
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)

                # Título principal
                doc.add_paragraph()  # Linha em branco
                doc.add_paragraph()  # Linha em branco
                title = doc.add_heading('MEMORIAL DESCRITIVO', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].font.name = 'Times New Roman'
                title.runs[0].font.size = Pt(16)
                title.runs[0].bold = True
                title.runs[0].underline = True
                doc.add_paragraph()  # Linha em branco
                doc.add_paragraph()  # Linha em branco

                # Seção 1: Beneficiário(s)
                heading_1 = doc.add_heading('1. Beneficiário(s):', level=1)
                heading_1.runs[0].font.name = 'Times New Roman'
                heading_1.runs[0].font.size = Pt(14)
                heading_1.runs[0].bold = True

                # Tabela de beneficiários
                table_ben = doc.add_table(rows=1 if not beneficiarios else len(beneficiarios), cols=2)
                table_ben.style = 'Table Grid'
                table_ben.autofit = False
                column_widths_ben = [Cm(8.0), Cm(8.0)]
                for col_idx, width in enumerate(column_widths_ben):
                    for cell in table_ben.columns[col_idx].cells:
                        tcPr = cell._tc.get_or_add_tcPr()
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:w'), str(int(width * 567)))
                        tcW.set(qn('w:type'), 'dxa')
                        tcPr.append(tcW)

                if beneficiarios:
                    for idx, ben in enumerate(beneficiarios):
                        row_cells = table_ben.rows[idx].cells
                        row_cells[0].text = ben.nome
                        row_cells[1].text = f"- CPF: {ben.cpf_cnpj}"
                        for cell in row_cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(12)
                            if cell == row_cells[0]:
                                cell.paragraphs[0].runs[0].bold = True
                else:
                    row_cells = table_ben.rows[0].cells
                    row_cells[0].text = "Nenhum beneficiário registrado."
                    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)

                doc.add_paragraph()
                doc.add_paragraph()

                # Seção 2: Localização do Imóvel
                heading_2 = doc.add_heading('2. Localização do Imóvel:', level=1)
                heading_2.runs[0].font.name = 'Times New Roman'
                heading_2.runs[0].font.size = Pt(14)
                heading_2.runs[0].bold = True

                p = doc.add_paragraph(f"{projeto.endereco}")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

                doc.add_paragraph()
                doc.add_paragraph()

                # Seção 3: Área
                heading_3 = doc.add_heading('3. Área:', level=1)
                heading_3.runs[0].font.name = 'Times New Roman'
                heading_3.runs[0].font.size = Pt(14)
                heading_3.runs[0].bold = True

                p = doc.add_paragraph(f"{projeto.area}m²")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

                doc.add_paragraph()
                doc.add_paragraph()

                # Seção 4: Perímetro
                heading_4 = doc.add_heading('4. Perímetro:', level=1)
                heading_4.runs[0].font.name = 'Times New Roman'
                heading_4.runs[0].font.size = Pt(14)
                heading_4.runs[0].bold = True

                p = doc.add_paragraph(f"{projeto.perimetro} m")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

                doc.add_paragraph()
                doc.add_paragraph()

                # Seção 5: Época da Medição
                heading_5 = doc.add_heading('5. Época da Medição:', level=1)
                heading_5.runs[0].font.name = 'Times New Roman'
                heading_5.runs[0].font.size = Pt(14)
                heading_5.runs[0].bold = True

                p = doc.add_paragraph(f"{projeto.epoca_medicao}")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

                doc.add_paragraph()
                doc.add_paragraph()

                # Seção 6: Instrumento Utilizado
                heading_6 = doc.add_heading('6. Instrumento Utilizado:', level=1)
                heading_6.runs[0].font.name = 'Times New Roman'
                heading_6.runs[0].font.size = Pt(14)
                heading_6.runs[0].bold = True

                p = doc.add_paragraph(f"{projeto.instrumento}")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

                doc.add_paragraph()
                doc.add_paragraph()

                # Seção 7: Sistema Geodésico de Referência
                heading_7 = doc.add_heading('7. Sistema Geodésico de Referência:', level=1)
                heading_7.runs[0].font.name = 'Times New Roman'
                heading_7.runs[0].font.size = Pt(14)
                heading_7.runs[0].bold = True

                p = doc.add_paragraph("SIRGAS 2000")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

                doc.add_paragraph()
                doc.add_paragraph()

                # Seção 8: Projeção Cartográfica de Distância e Área
                heading_8 = doc.add_heading('8. Projeção Cartográfica de Distância e Área:', level=1)
                heading_8.runs[0].font.name = 'Times New Roman'
                heading_8.runs[0].font.size = Pt(14)
                heading_8.runs[0].bold = True

                p = doc.add_paragraph("UTM")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

                doc.add_paragraph()
                doc.add_paragraph()

                # Seção 9: Tabela de Coordenadas, Confrontações e Medidas
                heading_9 = doc.add_heading('9. Tabela de Coordenadas, Confrontações e Medidas:', level=1)
                heading_9.runs[0].font.name = 'Times New Roman'
                heading_9.runs[0].font.size = Pt(14)
                heading_9.runs[0].bold = True

                # Criar tabela
                table = doc.add_table(rows=2 if not vertices else len(vertices) + 1, cols=6)
                table.style = 'Table Grid'
                table.autofit = False

                # Definir larguras das colunas
                column_widths = [Cm(2.0), Cm(1.5), Cm(3.5), Cm(3.5), Cm(2.0), Cm(3.5)]
                for col_idx, width in enumerate(column_widths):
                    for cell in table.columns[col_idx].cells:
                        tcPr = cell._tc.get_or_add_tcPr()
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:w'), str(int(width * 567)))
                        tcW.set(qn('w:type'), 'dxa')
                        tcPr.append(tcW)

                # Cabeçalho da tabela
                headers = ['DE', 'PARA', 'LONGITUDE', 'LATITUDE', 'DIST.', 'CONFRONTANTE']
                for i, header in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'D3D3D3')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                    cell.paragraphs[0].runs[0].bold = True

                # Preencher a tabela com os vértices ou mensagem de vazio
                if vertices:
                    for idx, ver in enumerate(vertices, start=1):
                        row_cells = table.rows[idx].cells
                        row_cells[0].text = str(ver.de_vertice)
                        row_cells[1].text = str(ver.para_vertice)
                        row_cells[2].text = str(ver.longitude)
                        row_cells[3].text = str(ver.latitude)
                        row_cells[4].text = f'{float(ver.distancia)} m'
                        row_cells[5].text = str(ver.confrontante.nome if ver.confrontante else ver.confrontante_texto)

                        for cell in row_cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                else:
                    row_cells = table.rows[1].cells
                    row_cells[0].text = "Nenhum vértice registrado."
                    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)

                doc.add_paragraph()
                doc.add_paragraph()

                # Local e Data
                p = doc.add_paragraph(f"{projeto.endereco.split(',')[-1].strip()}, 23 de Dezembro de 2024.")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

                doc.add_paragraph()

                # Assinatura do Responsável Técnico
                p = doc.add_paragraph("__________________________________________________")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

                p = doc.add_paragraph("Everton Valdir Pinto Vieira")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)
                p.runs[0].bold = True

                p = doc.add_paragraph("Resp. Técnico Técnico em Agrimensura")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

                p = doc.add_paragraph("CFT/CRT4025.441.619-57")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

                doc.add_paragraph()

                # Tabela de Assinaturas (Requerentes e Confrontantes)
                all_signatures = [(ben.nome, ben.cpf_cnpj, "Requerente") for ben in beneficiarios] + \
                                [(con.nome, con.cpf_cnpj, "Confrontante") for con in confrontantes]
                rows = max(1, (len(all_signatures) + 1) // 2)  # Pelo menos 1 linha
                table_sign = doc.add_table(rows=rows, cols=3)
                table_sign.style = None
                table_sign.autofit = False
                column_widths_sign = [Cm(8.0), Cm(1.0), Cm(7.0)]
                for col_idx, width in enumerate(column_widths_sign):
                    for cell in table_sign.columns[col_idx].cells:
                        tcPr = cell._tc.get_or_add_tcPr()
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:w'), str(int(width * 567)))
                        tcW.set(qn('w:type'), 'dxa')
                        tcPr.append(tcW)

                if all_signatures:
                    for idx, (nome, cpf_cnpj, tipo) in enumerate(all_signatures):
                        row_idx = idx // 2
                        col_idx = 0 if idx % 2 == 0 else 2
                        cell = table_sign.rows[row_idx].cells[col_idx]
                        p = cell.paragraphs[0]
                        run = p.add_run(f"{nome}\nCPF: {cpf_cnpj}\n{tipo}")
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        if nome in ["Alcides De Oliveira", "Maria Aparecida Trindade Oliveira"]:
                            run.underline = True
                else:
                    cell = table_sign.rows[0].cells[0]
                    p = cell.paragraphs[0]
                    run = p.add_run("Nenhuma assinatura registrada.")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                response = HttpResponse(
                    buffer.getvalue(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="memorial_descritivo_{projeto.nome}.docx"'
                return response
            except Projeto.DoesNotExist:
                messages.error(request, 'Projeto selecionado não existe.')
            except Exception as e:
                messages.error(request, f'Erro ao gerar memorial: {str(e)}')
                print(f"Erro detalhado: {str(e)}")

        elif action == 'gerar_memorial_pdf':
            projeto_id = request.POST.get('projeto_memorial')
            try:
                projeto = Projeto.objects.get(id=projeto_id)
                vertices = Vertice.objects.filter(projeto=projeto)
                beneficiarios = Beneficiario.objects.filter(projeto=projeto)
                confrontantes = Confrontante.objects.filter(projeto=projeto, excluir_do_pdf=False)



                # Log para depuração
                print(f"Projeto ID (PDF): {projeto_id}")
                print(f"Beneficiários encontrados (PDF): {len(beneficiarios)}")
                print(f"Confrontantes encontrados (PDF): {len(confrontantes)}")
                print(f"Vértices encontrados (PDF): {len(vertices)}")

                # Buffer para o PDF
                buffer = io.BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2.5*cm, leftMargin=2.5*cm, topMargin=2.5*cm, bottomMargin=1.5*cm)
                elements = []

                # Estilos
                styles = getSampleStyleSheet()
                title_style = ParagraphStyle(
                    'TitleStyle',
                    parent=styles['Heading1'],
                    fontName='Times-Roman',
                    fontSize=16,
                    alignment=1,  # Centro
                    spaceAfter=12,
                    textTransform='uppercase',
                    fontWeight='bold',
                    underline=True
                )
                heading_style = ParagraphStyle(
                    'HeadingStyle',
                    parent=styles['Heading2'],
                    fontName='Times-Roman',
                    fontSize=14,
                    spaceAfter=12,
                    fontWeight='bold'
                )
                normal_style = ParagraphStyle(
                    'NormalStyle',
                    parent=styles['Normal'],
                    fontName='Times-Roman',
                    fontSize=12,
                    spaceAfter=12,
                    firstLineIndent=1.25*cm,
                    alignment=4,  # Justificado
                    leading=5  # Espaçamento de 1,5 linhas (12pt * 1.5 = 18pt)
                )
                center_style = ParagraphStyle(
                    'CenterStyle',
                    parent=styles['Normal'],
                    fontName='Times-Roman',
                    fontSize=12,
                    spaceAfter=12,
                    alignment=1  # Centro
                )
                left_style = ParagraphStyle(
                    'LeftStyle',
                    parent=styles['Normal'],
                    fontName='Times-Roman',
                    fontSize=12,
                    spaceAfter=12,
                    alignment=0  # Left
                )

                # Título principal
                elements.append(Paragraph("MEMORIAL DESCRITIVO", title_style))
                elements.append(Paragraph("<br/><br/>", normal_style))  # Linhas em branco

                # Seção 1: Beneficiário(s)
                elements.append(Paragraph("1. Beneficiário(s):", heading_style))
                if beneficiarios:
                    data = []
                    for ben in beneficiarios:
                        data.append([Paragraph(ben.nome, ParagraphStyle('Bold', fontName='Times-Bold', fontSize=12)), f" -   CPF: {ben.cpf_cnpj}"])
                    table_ben = Table(data, colWidths=[8*cm, 8*cm])
                    table_ben.setStyle(TableStyle([
                        #('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
                        ('FONTSIZE', (0, 0), (-1, -1), 12),
                    ]))
                    elements.append(table_ben)
                else:
                    elements.append(Paragraph("Nenhum beneficiário registrado.", normal_style))
                elements.append(Paragraph("<br/>", normal_style))

                # Seção 2: Localização do Imóvel
                elements.append(Paragraph("2. Localização do Imóvel:", heading_style))
                elements.append(Paragraph(f"{projeto.endereco}", normal_style))
                elements.append(Paragraph("<br/>", normal_style))

                # Seção 3: Área
                elements.append(Paragraph("3. Área:", heading_style))
                elements.append(Paragraph(f"{projeto.area}m²", normal_style))
                elements.append(Paragraph("<br/>", normal_style))

                # Seção 4: Perímetro
                elements.append(Paragraph("4. Perímetro:", heading_style))
                elements.append(Paragraph(f"{projeto.perimetro} m", normal_style))
                elements.append(Paragraph("<br/>", normal_style))

                # Seção 5: Época da Medição
                elements.append(Paragraph("5. Época da Medição:", heading_style))
                elements.append(Paragraph(f"{projeto.epoca_medicao}", normal_style))
                elements.append(Paragraph("<br/>", normal_style))

                # Seção 6: Instrumento Utilizado
                elements.append(Paragraph("6. Instrumento Utilizado:", heading_style))
                elements.append(Paragraph(f"{projeto.instrumento}", normal_style))
                elements.append(Paragraph("<br/>", normal_style))

                # Seção 7: Sistema Geodésico de Referência
                elements.append(Paragraph("7. Sistema Geodésico de Referência:", heading_style))
                elements.append(Paragraph("SIRGAS 2000", normal_style))
                elements.append(Paragraph("<br/>", normal_style))

                # Seção 8: Projeção Cartográfica de Distância e Área
                elements.append(Paragraph("8. Projeção Cartográfica de Distância e Área:", heading_style))
                elements.append(Paragraph("UTM", normal_style))
                elements.append(Paragraph("<br/>", normal_style))

                # Seção 9: Tabela de Coordenadas, Confrontações e Medidas
                elements.append(Paragraph("9. Tabela de Coordenadas, Confrontações e Medidas:", heading_style))
                data = [['DE', 'PARA', 'LONGITUDE', 'LATITUDE', 'DIST.(m)', 'CONFRONTANTE']]
                if vertices:
                    for ver in vertices:
                        data.append([
                            str(ver.de_vertice),
                            str(ver.para_vertice),
                            str(ver.longitude),
                            str(ver.latitude),
                            f'{float(ver.distancia)} ',
                            str(ver.confrontante.nome if ver.confrontante else ver.confrontante_texto)
                        ])
                else:
                    data.append(["Nenhum vértice registrado.", "", "", "", "", ""])
                table = Table(data, colWidths=[1.5*cm, 1.5*cm, 3*cm, 3*cm, 1.5*cm, 4.5*cm])
                table.setStyle(TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTWEIGHT', (0, 0), (-1, 0), 'BOLD'),
                ]))
                elements.append(table)
                elements.append(Paragraph("<br/><br/>", normal_style))

                # Local e Data
                # Obter a data atual
                data_atual = datetime.now()

                # Dicionário para traduzir os meses para o português
                meses = {
                    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
                    7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
                }

                # Formatar a data no formato desejado (ex.: "28 de Abril de 2025")
                data_formatada = f"{data_atual.day} de {meses[data_atual.month]} de {data_atual.year}"

                # Adicionar o parágrafo com a data atual
                beneficiario = projeto.beneficiarios.first()
                cidade_beneficiario = beneficiario.cidade if beneficiario else "Cidade não especificada"
                elements.append(Paragraph(f"{cidade_beneficiario}, {data_formatada}.", left_style))
                elements.append(Paragraph("<br/>", normal_style))
                elements.append(Paragraph("<br/>", normal_style))
                elements.append(Paragraph("<br/>", normal_style))
                elements.append(Paragraph("<br/>", normal_style))
                elements.append(Paragraph("<br/>", normal_style))

                # Assinatura do Responsável Técnico
                elements.append(Paragraph("__________________________________________________", center_style))
                elements.append(Paragraph("Everton Valdir Pinto Vieira", ParagraphStyle('BoldCenter', parent=center_style, fontName='Times-Bold', fontWeight='bold')))
                elements.append(Paragraph("Resp. Técnico Técnico em Agrimensura", center_style))
                elements.append(Paragraph("CFT 4025.441.619-57", center_style))
                elements.append(Paragraph("<br/>", center_style))
                elements.append(Paragraph("<br/>", normal_style))
                elements.append(Paragraph("<br/>", normal_style))
                elements.append(Paragraph("<br/>", normal_style))
                elements.append(Paragraph("<br/>", normal_style))

                # Tabela de Assinaturas (Requerentes e Confrontantes)
                all_signatures = [(ben.nome, ben.cpf_cnpj, "Requerente") for ben in beneficiarios] + \
                                [(con.nome, con.cpf_cnpj, "Confrontante") for con in confrontantes]
                if all_signatures:
                    signature_data = []
                    for i in range(0, len(all_signatures), 2):
                        row = ["", "", ""]
                        # Primeira coluna
                        nome1, cpf1, tipo1 = all_signatures[i]
                        text1 = f"{nome1}<br/>CPF: {cpf1}<br/>{tipo1}"
                        style1 = ParagraphStyle('Signature', fontName='Times-Roman', fontSize=12, leading=14)
                        if nome1 in ["Alcides De Oliveira", "Maria Aparecida Trindade Oliveira"]:
                            text1 = f"<u>{nome1}</u><br/>CPF: {cpf1}<br/>{tipo1}"
                        row[0] = Paragraph(text1, style1)
                        # Segunda coluna (espaço)
                        row[1] = ""
                        # Terceira coluna (se houver)
                        if i + 1 < len(all_signatures):
                            nome2, cpf2, tipo2 = all_signatures[i + 1]
                            text2 = f"{nome2}<br/>CPF: {cpf2}<br/>{tipo2}"
                            style2 = ParagraphStyle('Signature', fontName='Times-Roman', fontSize=12, leading=14)
                            if nome2 in ["Alcides De Oliveira", "Maria Aparecida Trindade Oliveira"]:
                                text2 = f"<u>{nome2}</u><br/>CPF: {cpf2}<br/>{tipo2}"
                            row[2] = Paragraph(text2, style2)
                        signature_data.append(row)
                        # Adicionar duas linhas vazias após cada par de assinaturas
                        signature_data.append(["", "", ""])  # Primeira linha vazia
                        signature_data.append(["", "", ""])  # Segunda linha vazia
                        signature_data.append(["", "", ""])  # Segunda linha vazia

                    table_sign = Table(signature_data, colWidths=[8*cm, 1*cm, 7*cm])
                    table_sign.setStyle(TableStyle([
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ]))
                    elements.append(table_sign)
                else:
                    elements.append(Paragraph("Nenhuma assinatura registrada.", normal_style))

                # Gerar o PDF
                doc.build(elements)
                buffer.seek(0)

                response = HttpResponse(
                    buffer.getvalue(),
                    content_type='application/pdf'
                )
                response['Content-Disposition'] = f'attachment; filename="memorial_descritivo_{projeto.nome}.pdf"'
                return response
            except Projeto.DoesNotExist:
                messages.error(request, 'Projeto selecionado não existe.')
            except Exception as e:
                messages.error(request, f'Erro ao gerar memorial em PDF: {str(e)}')
                print(f"Erro detalhado (PDF): {str(e)}")

        if projeto_selecionado:
            beneficiarios = Beneficiario.objects.filter(projeto=projeto_selecionado)
            confrontantes = Confrontante.objects.filter(projeto=projeto_selecionado)
            vertices = Vertice.objects.filter(projeto=projeto_selecionado)

    return render(request, 'levantamento/index.html', {
        'projetos': projetos,
        'beneficiarios': beneficiarios,
        'confrontantes': confrontantes,
        'vertices': vertices,
        'projeto_selecionado': projeto_selecionado
    })


